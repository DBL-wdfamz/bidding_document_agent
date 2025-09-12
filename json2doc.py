# -*- coding: utf-8 -*-
"""
json2doc
两个对外函数：
  1) render_docx(json_data: dict, out_docx) -> Path
  2) docx_to_pdf(docx_path, pdf_path) -> bool
依赖：
  pip install python-docx
可选 PDF：
  - Windows: pip install docx2pdf
  - Linux/macOS: 安装 LibreOffice (soffice) 或 pandoc
"""
from __future__ import annotations
import json, os, shutil, subprocess
from pathlib import Path
from typing import Dict, Any, List, Optional, Union
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENTATION

PathLike = Union[str, Path]

# ------------------ 内部工具 ------------------

def _cm(v) -> Cm:
    if isinstance(v, (int, float)): return Cm(float(v))
    s = str(v).strip().lower().replace("cm","").strip()
    return Cm(float(s or 0))

def _pt(v) -> Pt:
    if isinstance(v, (int, float)): return Pt(float(v))
    s = str(v).strip().lower().replace("pt","").strip()
    return Pt(float(s or 0))

def _rgb(hexstr: str) -> RGBColor:
    s = str(hexstr).strip()
    if s.startswith("#"): s = s[1:]
    if len(s) == 6:
        return RGBColor(int(s[0:2],16), int(s[2:4],16), int(s[4:6],16))
    return RGBColor(0,0,0)

def _set_run_font(run, family: str):
    run.font.name = family
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), family)
    rFonts.set(qn('w:hAnsi'), family)
    rFonts.set(qn('w:eastAsia'), family)

def _style_hex_color(st: Dict[str,Any], fallback: str = "000000") -> str:
    col = (st or {}).get("color")
    if not col:
        return fallback
    s = str(col).strip()
    if s.startswith("#"): s = s[1:]
    return s.upper() if len(s) == 6 else fallback

def _apply_run_style(run, st: Dict[str, Any]):
    if not st: return
    if 'fontFamily' in st: _set_run_font(run, st['fontFamily'])
    if 'fontSize'  in st: run.font.size = _pt(st['fontSize'])
    if 'color'     in st: run.font.color.rgb = _rgb(st['color'])
    if st.get('bold'):   run.bold = True
    if st.get('italic'): run.italic = True

def _apply_para_style(p, st: Dict[str, Any]):
    if not st: return
    am = {'left':WD_ALIGN_PARAGRAPH.LEFT,'right':WD_ALIGN_PARAGRAPH.RIGHT,
          'center':WD_ALIGN_PARAGRAPH.CENTER,'justify':WD_ALIGN_PARAGRAPH.JUSTIFY}
    if 'alignment' in st:
        p.alignment = am.get(st['alignment'], WD_ALIGN_PARAGRAPH.LEFT)
    sp = st.get('spacing') or {}
    if 'before' in sp: p.paragraph_format.space_before = _pt(sp['before'])
    if 'after'  in sp: p.paragraph_format.space_after  = _pt(sp['after'])
    if 'line'   in sp: p.paragraph_format.line_spacing = _pt(sp['line'])
    ind = st.get('indentation') or {}
    if 'firstLine' in ind:
        val = str(ind['firstLine']).strip().lower()
        if val.endswith('em'):
            try: em = float(val[:-2])
            except: em = 2.0
            p.paragraph_format.first_line_indent = Cm(0.74*em)  # 近似中文排版
        else:
            p.paragraph_format.first_line_indent = _cm(val)
    if 'right' in ind:
        p.paragraph_format.right_indent = _cm(ind['right'])

def _add_header(section, text: str, st: Dict[str,Any], red_line=False):
    header = section.header
    p = header.add_paragraph()
    _apply_para_style(p, st or {})
    r = p.add_run(text or "")
    _apply_run_style(r, st or {})
    if red_line:
        pPr = p._p.get_or_add_pPr()
        pBdr = pPr.find(qn('w:pBdr')) or OxmlElement('w:pBdr')
        if pBdr not in list(pPr): pPr.append(pBdr)
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),'12'); bottom.set(qn('w:space'),'1'); bottom.set(qn('w:color'),'FF0000')
        pBdr.append(bottom)

def _add_footer_page_number(section, st: Optional[Dict[str,Any]]=None):
    p = section.footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run("- "); _apply_run_style(r1, st or {})
    # 插入 PAGE 字段
    r = OxmlElement('w:r'); fldBegin = OxmlElement('w:fldChar'); fldBegin.set(qn('w:fldCharType'),'begin'); r.append(fldBegin); p._p.append(r)
    r = OxmlElement('w:r'); instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'),'preserve'); instr.text = ' PAGE '; r.append(instr); p._p.append(r)
    r = OxmlElement('w:r'); fldEnd = OxmlElement('w:fldChar');  fldEnd.set(qn('w:fldCharType'),'end');   r.append(fldEnd);  p._p.append(r)
    r3 = p.add_run(" -"); _apply_run_style(r3, st or {})

def _apply_page(doc: Document, page_setup: Dict[str,Any], styles: Dict[str,Any]):
    sec = doc.sections[0]
    # 页面尺寸与方向
    if str(page_setup.get('pageSize','A4')).upper() == 'A4':
        sec.page_height, sec.page_width = Cm(29.7), Cm(21.0)
    if str(page_setup.get('orientation','portrait')).lower() == 'landscape':
        sec.orientation = WD_ORIENTATION.LANDSCAPE
        sec.page_width, sec.page_height = sec.page_height, sec.page_width
    # 边距
    m = page_setup.get('margins') or {}
    if 'top' in m:    sec.top_margin    = _cm(m['top'])
    if 'bottom' in m: sec.bottom_margin = _cm(m['bottom'])
    if 'left' in m:   sec.left_margin   = _cm(m['left'])
    if 'right' in m:  sec.right_margin  = _cm(m['right'])

    hd = m.get('yemei_top') or m.get('headerTop') or m.get('header_distance')
    if hd is not None:
        sec.header_distance = _cm(hd)

    fd = (m.get('yejiao_bottom') or m.get('footerBottom') or
          m.get('footer_distance') or m.get('yemei_bottom'))  # 兼容把 yemei_bottom 当“页脚距边界”的旧数据
    if fd is not None:
        sec.footer_distance = _cm(fd)

    # --------- 页眉：首页与非首页 ----------
    hongtou = page_setup.get('redHeader')  # {"style":"hongtou", "content":"xxx", "separatorLine":true}
    header_cfg = (page_setup.get('yemei') or {}).get('content', [])  # 列表，每项含 style、content

    if hongtou:
        sec.different_first_page_header_footer = True
        fp_hdr = sec.first_page_header

        # 1) 红头（仅文字，不加红线）
        st = styles.get(hongtou.get('style', ''), {})
        p_head = fp_hdr.add_paragraph()
        _apply_para_style(p_head, st)
        r_head = p_head.add_run(str(hongtou.get('content', '')))
        _apply_run_style(r_head, st)

        # 2) 文号（紧跟红头）
        wenhao = page_setup.get('wenhao')
        p_line_target = None  # 准备给“红线”使用的目标段落
        if isinstance(wenhao, dict):
            wst = styles.get(wenhao.get('style', ''), {})
            p_wenhao = fp_hdr.add_paragraph()
            _apply_para_style(p_wenhao, wst)
            r_wenhao = p_wenhao.add_run(str(wenhao.get('content', '')))
            _apply_run_style(r_wenhao, wst)
            p_line_target = p_wenhao
        else:
            # 没有 wenhao，则仍然把红线画在红头下（兼容旧数据）
            p_line_target = p_head

        # 3) 红线：现在画在“文号”段落的下边框（若配置了 separatorLine）
        if str(hongtou.get('separatorLine', '')).lower() in ('true', '1', 'yes'):
            # 线条颜色优先取红头样式，如果没有则默认 #FF0000
            _maybe_add_bottom_red_line(p_line_target, _style_hex_color(st, "FF0000"))

        # 其余页面眉：用 header.content
        for item in header_cfg:
            stx = styles.get(item.get('style', ''), {})
            xflag = str(item.get('separatorLine', '')).lower() in ('true', '1', 'yes')
            _add_text_with_optional_page_field(sec.header, item.get('content', ''), stx, xflag)
    else:
        # 没有 redHeader，则所有页都用 header.content
        for item in header_cfg:
            stx = styles.get(item.get('style', ''), {})
            xflag = str(item.get('separatorLine', '')).lower() in ('true', '1', 'yes')
            _add_text_with_optional_page_field(sec.header, item.get('content', ''), stx, xflag)

    # --------- 页脚：全部页面一致 ----------
    def _render_footer_to(container, items, styles):
        # 先渲染 footer.content
        for it in (items or []):
            stx = styles.get(it.get('style',''), {})
            _add_text_with_optional_page_field(container, it.get('content',''), stx)

        # 再渲染 pageNumber（若提供）
        pn = page_setup.get('pageNumber')
        if pn:
            stx = styles.get(pn.get('style',''), {})
            fmt = pn.get('format', '')
            _add_text_with_optional_page_field(container, fmt, stx)

    footer_cfg = (page_setup.get('footer') or {}).get('content', [])
    # 普通页页脚
    _render_footer_to(sec.footer, footer_cfg, styles)

    # 如果启用了不同首页，单独渲染第一页页脚
    if getattr(sec, 'different_first_page_header_footer', False):
        _render_footer_to(sec.first_page_footer, footer_cfg, styles)

def _add_paragraph_with_runs(doc: Document, pieces: List[Dict[str,Any]], base_style: Dict[str,Any], styles: Dict[str,Any]):
    p = doc.add_paragraph()
    _apply_para_style(p, base_style or {})
    for piece in pieces:
        text = piece.get('text','')
        piece_style = styles.get(piece.get('style',''), base_style)
        r = p.add_run(text); _apply_run_style(r, piece_style or {})
    return p

def _add_text_with_optional_page_field(container, text: str, st: Dict[str,Any], xian: bool=False):
    """
    在 header/footer 的段落里渲染文本；若包含 %p 则插入 PAGE 字段。
    container: section.header / section.first_page_header / section.footer
    """
    p = container.add_paragraph()
    _apply_para_style(p, st or {})

    # 新增：若需要画红线
    if xian:
        _maybe_add_bottom_red_line(p, _style_hex_color(st, "000000"))

    if text is None:
        return

    s = str(text)
    if "%p" not in s:
        r = p.add_run(s)
        _apply_run_style(r, st or {})
        return

    # 拆分前后缀
    prefix, suffix = s.split("%p", 1)
    if prefix:
        r1 = p.add_run(prefix)
        _apply_run_style(r1, st or {})

    # 插入 Word 的 PAGE 字段
    r = OxmlElement('w:r'); fldBegin = OxmlElement('w:fldChar'); fldBegin.set(qn('w:fldCharType'),'begin'); r.append(fldBegin); p._p.append(r)
    r = OxmlElement('w:r'); instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'),'preserve'); instr.text = ' PAGE '; r.append(instr); p._p.append(r)
    r = OxmlElement('w:r'); fldEnd = OxmlElement('w:fldChar');  fldEnd.set(qn('w:fldCharType'),'end');   r.append(fldEnd);  p._p.append(r)

    if suffix:
        r3 = p.add_run(suffix)
        _apply_run_style(r3, st or {})


def _maybe_add_bottom_red_line(paragraph, color_hex: Optional[str] = None):
    """
    为段落添加底部细线；颜色默认取调用方传入的 hex（不含# 或含都可）。
    """
    hexv = (color_hex or "FF0000").strip()
    if hexv.startswith("#"): hexv = hexv[1:]
    if len(hexv) != 6: hexv = "FF0000"

    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr')) or OxmlElement('w:pBdr')
    if pBdr not in list(pPr):
        pPr.append(pBdr)
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),'single')
    bottom.set(qn('w:sz'),'12')
    bottom.set(qn('w:space'),'1')
    bottom.set(qn('w:color'), hexv)
    pBdr.append(bottom)

def _render_block(doc: Document, block: Dict[str,Any], styles: Dict[str,Any]):
    """
    不再使用 blockType；所有块都渲染为段落，按 block['style'] 取段落样式。
    若 block['content'] 是列表（每项可带 text/style），将逐个 run 渲染；
    若是字符串，则整段一个 run。
    """
    st = styles.get(block.get("style",""), {})

    # 1) content 列表：按子项 style 优先
    if isinstance(block.get("content"), list):
        _add_paragraph_with_runs(doc, block["content"], st, styles)
        return

    # 2) 单文本：content 或 text
    txt = block.get("content")
    if txt is None:
        txt = block.get("text", "")
    p = doc.add_paragraph()
    _apply_para_style(p, st)
    r = p.add_run(str(txt))
    _apply_run_style(r, st)

def _which(cmd: str) -> Optional[str]:
    return shutil.which(cmd)

# ------------------ 对外函数 1：渲染 DOCX ------------------

def render_docx(json_data: Dict[str, Any], out_docx: PathLike) -> Path:
    """
    将符合约定结构的 JSON（dict）渲染为 DOCX。
    参数：
      json_data: dict（示例结构与“模板定义”一致）
      out_docx:  输出 DOCX 路径
    返回：
      实际写入的 Path
    """
    outp = Path(out_docx)
    outp.parent.mkdir(parents=True, exist_ok=True)

    styles = json_data.get("styles", {}) or {}
    page_setup = json_data.get("pageSetup", {}) or {}
    blocks = json_data.get("contentBlocks", []) or []

    doc = Document()
    _apply_page(doc, page_setup, styles)
    for blk in blocks:
        _render_block(doc, blk, styles)

    doc.save(outp)
    return outp

# ------------------ 对外函数 2：DOCX 转 PDF ------------------

def docx_to_pdf(docx_path: PathLike, pdf_path: PathLike) -> bool:
    """
    将 DOCX 转为 PDF。
    优先级：
      1) Windows: docx2pdf
      2) LibreOffice/soffice
      3) pandoc
    成功返回 True，失败 False。
    """
    docx_path = Path(docx_path)
    pdf_path = Path(pdf_path)
    pdf_path.parent.mkdir(parents=True, exist_ok=True)

    # 1) Windows + docx2pdf
    if os.name == "nt":
        try:
            from docx2pdf import convert as _convert
            _convert(str(docx_path), str(pdf_path))
            return True
        except Exception:
            pass

    # 2) LibreOffice / soffice
    soffice = _which("soffice") or _which("libreoffice")
    if soffice:
        try:
            cmd = [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(pdf_path.parent), str(docx_path)]
            subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            generated = docx_path.with_suffix(".pdf")
            if generated.exists():
                if generated.resolve() != pdf_path.resolve():
                    generated.replace(pdf_path)
                return True
        except Exception:
            pass

    # 3) pandoc
    pandoc = _which("pandoc")
    if pandoc:
        try:
            subprocess.run([pandoc, str(docx_path), "-o", str(pdf_path)], check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            return True
        except Exception:
            pass

    return False

json_data = {
  "docInfo": {
    "docType": "招标公告",
    "title": "关于XX市智慧城市大脑平台（一期）建设项目的公开招标公告",
    "urgency": "平件",
    "security": "公开"
  },
  "pageSetup": {
    "pageSize": "A4",
    "orientation": "portrait",
    "margins": {
      "top": "3.7cm",
      "bottom": "3.5cm",
      "left": "2.8cm",
      "right": "2.6cm",
      "yemei_top": "1cm",
      "yejiao_bottom": "1cm"
    },
    "redHeader":{
      "style":"redHeader",
      "content":"某某公司文件",
      "separatorLine":"true"
    },
    "yemei": {
      "content": [
        { "style": "yemei", "content": "XX市公共资源交易中心文件", "separatorLine": "true" }
      ]
    },
    "footer": {
      "content": [
        { "style": "footerText", "content": "XX市公共资源交易中心文件" }
      ]
    },
    "pageNumber": {
      "style": "footerText",
      "format": "- %p -"
    },
    "wenhao": {
      "style": "wenhao",
      "content": "测试公司(2025)第107号"
    }
  },
  "styles": {
    "wenhao": { "fontFamily": "楷体_GB2312", "fontSize": "14pt", "alignment": "center", "spacing": { "after": "24pt" } },
    "yemei": { "fontFamily": "方正小标宋_GBK", "fontSize": "12pt", "color": "#000000", "spacing": { "after": "12pt" } },
    "mainTitle": { "fontFamily": "黑体", "fontSize": "22pt", "color": "#000000", "bold": "true", "alignment": "center", "spacing": { "before": "12pt", "after": "12pt" } },
    "docNumberText": { "fontFamily": "楷体_GB2312", "fontSize": "16pt", "alignment": "center", "spacing": { "after": "24pt" } },
    "heading1": { "fontFamily": "黑体", "fontSize": "16pt", "bold": "true", "spacing": { "before": "12pt", "after": "6pt" } },
    "heading2": { "fontFamily": "楷体_GB2312", "fontSize": "16pt", "bold": "true", "indentation": { "firstLine": "2em" }, "spacing": { "before": "6pt", "after": "6pt" } },
    "heading3": { "fontFamily": "仿宋_GB2312", "fontSize": "16pt", "bold": "true", "indentation": { "firstLine": "2em" }, "spacing": { "before": "6pt", "after": "6pt" } },
    "bodyText": { "fontFamily": "仿宋_GB2312", "fontSize": "16pt", "alignment": "justify", "indentation": { "firstLine": "2em" }, "spacing": { "line": "28.8pt", "lineRule": "exactly" } },
    "bodyTextBold": { "fontFamily": "仿宋_GB2312", "fontSize": "16pt", "bold": "true" },
    "tableText": { "fontFamily": "仿宋_GB2312", "fontSize": "14pt", "alignment": "center" },
    "signatureText": { "fontFamily": "仿宋_GB2312", "fontSize": "16pt", "alignment": "right", "spacing": { "before": "24pt", "line": "32pt", "lineRule": "exactly" }, "indentation": { "right": "5cm" } },
    "footerText": { "fontFamily": "宋体", "fontSize": "12pt", "alignment": "center" },
    "redHeader":{ "fontFamily": "方正小标宋_GBK", "fontSize": "25pt", "color": "#FF0000", "alignment": "center", "spacing": { "after": "12pt" } }
  },
  "contentBlocks": [
    {
      "style": "mainTitle",
      "content": "关于XX市智慧城市大脑平台（一期）建设项目的公开招标公告"
    },
    {
      "style": "docNumberText",
      "content": "招标编号：GXZB-2025-0828"
    },
    {
      "style": "heading1",
      "level": 1,
      "content": "项目概况"
    },
    {
      "style": "bodyText",
      "content": "XX市智慧城市大脑平台（一期）建设项目招标项目的潜在投标人应在XX市公共资源交易网获取招标文件，并于2025年9月20日09点30分（北京时间）前递交投标文件。"
    },
    {
      "style": "heading1",
      "content": "申请人的资格要求"
    },
    {
      "style": "heading2",
      "content": "满足《中华人民共和国政府采购法》第二十二条规定"
    },
    {
      "style": "heading3",
      "content": "具有独立承担民事责任的能力。"
    },
    {
      "style": "heading3",
      "content": "具有良好的商业信誉和健全的财务会计制度。"
    },
    {
      "style": "heading2",
      "content": "本项目的特定资格要求"
    },
    {
      "style": "bodyText",
      "content": "投标人须具备电子与智能化工程专业承包\n壹级\n及以上资质。"
    },
    {
      "style": "signatureText",
      "content": "XX招标代理有限公司\n二〇二五年八月二十八日"
    },
    {
      "style": "bodyText",
      "content": "附件1：项目技术需求书.pdf\n附件2：投标人资格证明文件模板.docx\n附件2：投标人资格证明文件模板.docx\n附件2：投标人资格证明文件模板.docx\n附件2：投标人资格证明文件模板.docx\n附件2：投标人资格证明文件模板.docx\n附件2：投标人资格证明文件模板.docx\n附件2：投标人资格证明文件模板.docx\n附件2：投标人资格证明文件模板.docx"
    }
  ]
}



# 1. 直接生成 DOCX
#docx_file = render_docx(json_data, "C:\\Users\\32271\\Desktop\\xxx.docx")
# 2. 如果需要 PDF
#ok = docx_to_pdf("D:\\project\\document_agent\\exports\\xxx.docx", "D:\\project\\document_agent\\exports\\xxx.pdf")
