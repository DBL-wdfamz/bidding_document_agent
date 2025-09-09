# -*- coding: utf-8 -*-
# file: knowledge_base.py

import os
import pandas as pd
from docx import Document
from typing import List, Set, Optional, Dict

# ==================== 模块零：全局配置与依赖 ====================
from huggingface_hub import snapshot_download
from sentence_transformers import CrossEncoder
from langchain.schema import Document as LangChainDocument
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS

class KnowledgeBase:
    """
    一个自包含的、可持久化的知识库模块。
    采用“召回+重排序”二级检索策略，以实现高精度问答。
    """

    # ==================== 模块一：初始化 ====================
    def __init__(self, 
                 embedding_model_name: str = "BAAI/bge-large-zh-v1.5", 
                 reranker_model_name: str = "BAAI/bge-reranker-large",
                 db_path: Optional[str] = None):
        print("Initializing KnowledgeBase with Reranking support...")
        
        embedding_model_path = self._prepare_model(embedding_model_name)
        print(f"Loading embedding model from: {embedding_model_path}")
        self.embeddings = HuggingFaceEmbeddings(model_name=embedding_model_path)
        
        reranker_model_path = self._prepare_model(reranker_model_name)
        print(f"Loading reranker model from: {reranker_model_path}")
        self.reranker = CrossEncoder(reranker_model_path)
        
        self.db_path = db_path
        self.vectorstore = self._load_or_create_vectorstore()
        print("KnowledgeBase initialized successfully.")

    @staticmethod
    def _prepare_model(model_name: str) -> str:
        try:
            model_path = snapshot_download(model_name, local_files_only=True)
            print(f"Model '{model_name}' found in local cache.")
            return model_path
        except FileNotFoundError:
            print(f"Model '{model_name}' not found locally. Starting download...")
            model_path = snapshot_download(model_name, local_files_only=False)
            print("Model downloaded successfully.")
            return model_path

    def _load_or_create_vectorstore(self) -> FAISS:
        """从本地加载向量数据库，如果不存在则创建一个新的并立即持久化。"""
        if self.db_path and os.path.exists(self.db_path):
            try:
                print(f"Loading vectorstore from {self.db_path}")
                return FAISS.load_local(self.db_path, self.embeddings, allow_dangerous_deserialization=True)
            except Exception as e:
                print(f"Error loading vectorstore: {e}. Re-creating a new one.")
        
        # <<< 核心修复点：创建新知识库时，立即在磁盘上创建文件夹和文件 >>>
        print("Creating a new vectorstore and persisting it immediately.")
        # 1. 创建一个空的向量存储
        dummy_doc = [LangChainDocument(page_content="init")]
        new_vectorstore = FAISS.from_documents(dummy_doc, self.embeddings)
        
        # 2. 立即将其保存到磁盘，这会自动创建所有必要的父文件夹
        if self.db_path:
            new_vectorstore.save_local(self.db_path)
            print(f"New knowledge base created and saved at: {self.db_path}")
        
        return new_vectorstore

    # ==================== 模块二：数据处理与存储 ====================
    def save_local(self):
        """将向量数据库持久化到本地。"""
        if self.db_path and self.vectorstore:
            self.vectorstore.save_local(self.db_path)

    def add_documents(self, documents: List[LangChainDocument], chunk_size: int = 1000, chunk_overlap: int = 100):
        """将LangChain文档对象添加到知识库中。"""
        if not documents: return
        splits = self._split_documents(documents, chunk_size, chunk_overlap)
        if splits:
            self.vectorstore.add_documents(splits)
            self.save_local()

    # ==================== 模块三：数据检索 (核心优化) ====================
    def query_with_score(self, query_text: str, top_k: int = 4, recall_k: int = 20) -> List[Dict[str, any]]:
        """使用“召回+重排序”策略查询知识库，返回最相关的文档及最终的精确相关性得分。"""
        if not self.vectorstore: raise ValueError("Vectorstore is not initialized.")
        
        recalled_docs = self.vectorstore.similarity_search(query_text, k=recall_k)
        if not recalled_docs: return []

        sentence_pairs = [[query_text, doc.page_content] for doc in recalled_docs]
        rerank_scores = self.reranker.predict(sentence_pairs)

        reranked_results = []
        for doc, score in zip(recalled_docs, rerank_scores):
            reranked_results.append({"document": doc, "score": float(score)})
        
        reranked_results.sort(key=lambda x: x['score'], reverse=True)
        
        return reranked_results[:top_k]

    # ==================== 模块四：知识库管理 ====================
    def list_sources(self) -> Set[str]:
        """列出知识库中所有文档的来源。"""
        if not hasattr(self.vectorstore, 'docstore'): return set()
        sources = {doc.metadata.get('source') for doc in self.vectorstore.docstore._dict.values() if doc.metadata.get('source')}
        sources.discard("init")
        return sources

    def delete_documents_by_source(self, source_filename: str) -> bool:
        """根据来源文件名删除文档。"""
        if not hasattr(self.vectorstore, 'docstore'): return False
        ids_to_delete = [doc_id for doc_id, doc in self.vectorstore.docstore._dict.items() if doc.metadata.get('source') == source_filename]
        if not ids_to_delete: return False
        try:
            self.vectorstore.delete(ids_to_delete)
            self.save_local()
            return True
        except Exception:
            return False

    # ==================== 模块五：静态辅助工具 ====================
    @staticmethod
    def load_documents_from_paths(file_paths: List[str]) -> List[LangChainDocument]:
        """从文件路径列表加载文档，并统一元数据。"""
        final_docs = []
        for file_path in file_paths:
            file_name = os.path.basename(file_path)
            file_ext = os.path.splitext(file_name)[1].lower()
            try:
                loaded_docs = []
                if file_ext == ".pdf": loader = PyPDFLoader(file_path); loaded_docs = loader.load()
                elif file_ext in [".txt", ".md"]: loader = TextLoader(file_path, encoding='utf-8'); loaded_docs = loader.load()
                elif file_ext == ".docx":
                    doc = Document(file_path)
                    content = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                    if content: loaded_docs.append(LangChainDocument(page_content=content, metadata={"source": file_name}))
                elif file_ext in [".csv", ".xlsx"]:
                    df = pd.read_excel(file_path) if file_ext == ".xlsx" else pd.read_csv(file_path, encoding='utf-8')
                    for index, row in df.iterrows():
                        row_content = ", ".join([f"{col}: {val}" for col, val in row.items() if pd.notna(val)])
                        if row_content: loaded_docs.append(LangChainDocument(page_content=row_content, metadata={"source": file_name, "row": index + 1}))
                
                for doc in loaded_docs: doc.metadata['source'] = file_name
                final_docs.extend(loaded_docs)
            except Exception as e:
                print(f"Error loading {file_name}: {e}")
        return final_docs

    @staticmethod
    def _split_documents(documents: List[LangChainDocument], chunk_size: int, chunk_overlap: int) -> List[LangChainDocument]:
        """内部辅助方法：切分文档。"""
        return RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap).split_documents(documents)